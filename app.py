from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file
from datetime import datetime
import calendar
import io
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

app = Flask(__name__)
app.secret_key = "your_secret_key_here"  # REPLACE with a strong secret key!
app.jinja_env.globals.update(enumerate=enumerate)

# *********************************************************************
# IMPORTANT:
# 1. Ensure that every “Back” button in your HTML templates includes 
#    formnovalidate so that the browser does not enforce required-field validation.
#
# 2. In your activities template, use event delegation for dynamic quarter dropdowns.
# *********************************************************************

# -------------------------
# Session Initialization
# -------------------------
def init_session():
    if 'general_info' not in session:
        session['general_info'] = {}
    if 'outcomes' not in session:
        session['outcomes'] = []          # List of outcome dictionaries
    if 'outputs' not in session:
        session['outputs'] = []           # List of output dictionaries
    if 'output_budgets' not in session:
        session['output_budgets'] = {}    # Dict: key = output index (as str), value = list of budget entries
    if 'activities' not in session:
        session['activities'] = {}        # Dict: key = output index (as str), value = list of activity dictionaries
    if 'goods' not in session:
        session['goods'] = []             # List of goods procurement dictionaries
    if 'services' not in session:
        session['services'] = []          # List of services procurement dictionaries
    if 'use_goods' not in session:
        session['use_goods'] = False
    if 'use_services' not in session:
        session['use_services'] = False

# -------------------------
# Data Classes
# -------------------------
class Outcome:
    def __init__(self, description, pi, baseline, target, mov, assumptions):
        self.description = description
        self.pi = pi
        self.baseline = baseline
        self.target = target
        self.mov = mov
        self.assumptions = assumptions

    def to_dict(self):
        return {
            'description': self.description,
            'pi': self.pi,
            'baseline': self.baseline,
            'target': self.target,
            'mov': self.mov,
            'assumptions': self.assumptions
        }

    @classmethod
    def from_dict(cls, data):
        return cls(data['description'], data['pi'], data['baseline'],
                   data['target'], data['mov'], data['assumptions'])

class Output:
    def __init__(self, outcome_index, description, pi, baseline, target, mov, assumptions, number):
        self.outcome_index = outcome_index  # zero-based index of associated outcome
        self.description = description
        self.pi = pi
        self.baseline = baseline
        self.target = target
        self.mov = mov
        self.assumptions = assumptions
        self.number = number  # e.g., "1.1", "1.2", etc.

    def to_dict(self):
        return {
            'outcome_index': self.outcome_index,
            'description': self.description,
            'pi': self.pi,
            'baseline': self.baseline,
            'target': self.target,
            'mov': self.mov,
            'assumptions': self.assumptions,
            'number': self.number
        }

    @classmethod
    def from_dict(cls, data):
        return cls(data['outcome_index'], data['description'], data['pi'],
                   data['baseline'], data['target'], data['mov'], data['assumptions'], data['number'])

class Activity:
    def __init__(self, output_index, description, start_year, start_quarter, end_year, end_quarter):
        self.output_index = output_index
        self.description = description
        self.start_year = start_year
        self.start_quarter = start_quarter
        self.end_year = end_year
        self.end_quarter = end_quarter

    def to_dict(self):
        return {
            'output_index': self.output_index,
            'description': self.description,
            'start_year': self.start_year,
            'start_quarter': self.start_quarter,
            'end_year': self.end_year,
            'end_quarter': self.end_quarter
        }

    @classmethod
    def from_dict(cls, data):
        return cls(data['output_index'], data['description'], data['start_year'],
                   data['start_quarter'], data['end_year'], data['end_quarter'])

# -------------------------
# Helper Functions
# -------------------------
def quarter_of_month(m):
    if 1 <= m <= 3:
        return 1
    elif 4 <= m <= 6:
        return 2
    elif 7 <= m <= 9:
        return 3
    else:
        return 4

def get_quarter_number(q):
    q = q.strip().upper()
    if q == 'Q1': return 1
    if q == 'Q2': return 2
    if q == 'Q3': return 3
    if q == 'Q4': return 4
    return 0

def parse_date_string(date_str):
    return datetime.strptime(date_str, "%Y-%m-%d")

def FormatWithHardSpace(value, fmt=",.0f"):
    try:
        s = format(value, fmt)
    except Exception:
        s = str(value)
    return s.replace(",", "\u00A0")

def get_years_and_quarters(start_date, end_date):
    results = []
    current_year = start_date.year
    start_q = quarter_of_month(start_date.month)
    end_q = quarter_of_month(end_date.month)
    while current_year <= end_date.year:
        if current_year == start_date.year and current_year == end_date.year:
            quarters = [f"Q{qn}" for qn in range(start_q, end_q + 1)]
            results.append({"Year": current_year, "Quarters": quarters})
            break
        elif current_year == start_date.year:
            quarters = [f"Q{qn}" for qn in range(start_q, 5)]
            results.append({"Year": current_year, "Quarters": quarters})
        elif current_year == end_date.year:
            quarters = [f"Q{qn}" for qn in range(1, end_q + 1)]
            results.append({"Year": current_year, "Quarters": quarters})
        else:
            results.append({"Year": current_year, "Quarters": ["Q1", "Q2", "Q3", "Q4"]})
        current_year += 1
    return results

def interpolate_color(start_rgb, end_rgb, fraction):
    r = int(start_rgb[0] + (end_rgb[0] - start_rgb[0]) * fraction)
    g = int(start_rgb[1] + (end_rgb[1] - start_rgb[1]) * fraction)
    b = int(start_rgb[2] + (end_rgb[2] - start_rgb[2]) * fraction)
    return f"{r:02X}{g:02X}{b:02X}"

# -------------------------
# Cascade Deletion and Recalculation
# -------------------------
def recalc_output_numbers():
    outputs = session.get('outputs', [])
    grouped = {}
    for op in outputs:
        idx = op['outcome_index']
        grouped.setdefault(idx, []).append(op)
    for idx, ops in grouped.items():
        ops.sort(key=lambda x: float(x.get('number', '0')))
        count = 1
        for op in ops:
            op['number'] = f"{idx+1}.{count}"
            count += 1
    session['outputs'] = outputs

def delete_outcome_cascade(outcome_index):
    outcomes = session.get('outcomes', [])
    if 0 <= outcome_index < len(outcomes):
        outcomes.pop(outcome_index)
    outputs = session.get('outputs', [])
    new_outputs = []
    for op in outputs:
        if op['outcome_index'] == outcome_index:
            continue
        elif op['outcome_index'] > outcome_index:
            op['outcome_index'] -= 1
            new_outputs.append(op)
        else:
            new_outputs.append(op)
    session['outcomes'] = outcomes
    session['outputs'] = new_outputs
    old_budgets = session.get('output_budgets', {})
    old_acts = session.get('activities', {})
    new_budgets = {}
    new_acts = {}
    for new_idx in range(len(new_outputs)):
        key = str(new_idx)
        new_budgets[key] = old_budgets.get(key, [])
        new_acts[key] = old_acts.get(key, [])
    session['output_budgets'] = new_budgets
    session['activities'] = new_acts
    recalc_output_numbers()

def delete_output_cascade(output_index):
    outputs = session.get('outputs', [])
    new_outputs = []
    for i, op in enumerate(outputs):
        if i == output_index:
            continue
        new_outputs.append(op)
    session['outputs'] = new_outputs
    old_budgets = session.get('output_budgets', {})
    old_acts = session.get('activities', {})
    new_budgets = {}
    new_acts = {}
    for new_idx in range(len(new_outputs)):
        key = str(new_idx)
        new_budgets[key] = old_budgets.get(str(new_idx), [])
        new_acts[key] = old_acts.get(str(new_idx), [])
    session['output_budgets'] = new_budgets
    session['activities'] = new_acts
    recalc_output_numbers()

# -------------------------
# Word Document Generation Functions
# -------------------------
def generate_word_doc():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(16.54)
    section.page_height = Inches(11.69)

    gi = session.get('general_info', {})
    project_symbol = gi.get('project_symbol', 'N/A')
    operational_partner = gi.get('operational_partner', 'N/A')
    project_title = gi.get('project_title', 'N/A')
    opa_val = gi.get('opa_value', 0.0)

    # Header: dark, bold, left-aligned
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for label, value in [("Project Symbol: ", project_symbol),
                         ("Operational Partner: ", operational_partner),
                         ("Project Title: ", project_title),
                         ("OPA Value: ", FormatWithHardSpace(opa_val) + "\u00A0USD")]:
        run = header_para.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        header_para.add_run(f"{value}\n")
    doc.add_paragraph()

    generate_results_matrix(doc)
    doc.add_paragraph()
    generate_workplan(doc)
    doc.add_paragraph()
    generate_budget_table(doc)
    doc.add_paragraph()
    generate_budget_fao_table(doc)
    doc.add_paragraph()
    if session.get('use_goods', False) and len(session.get('goods', [])) > 0:
        generate_goods_table(doc)
        doc.add_paragraph()
    if session.get('use_services', False) and len(session.get('services', [])) > 0:
        generate_subcontracting_table(doc)
        doc.add_paragraph()
    return doc

def generate_results_matrix(doc):
    heading = doc.add_heading("RESULTS MATRIX", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    
    # Set header row properties and shading with explicit namespace
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(parse_xml('<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="true"/>'))
    hdr_cells = table.rows[0].cells
    headers = ["Result Statement", "Performance Indicators", "Baseline", "Target", "Means of Verification", "Assumptions"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
        shading = parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="CCE5FF"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    
    outcomes = session.get('outcomes', [])
    outputs = session.get('outputs', [])
    
    # For each outcome, add an outcome row then its outputs
    for idx, oc in enumerate(outcomes):
        # Add outcome row
        outcome_cells = table.add_row().cells
        outcome_cells[0].text = f"Outcome {idx+1}: {oc.get('description', '')}"
        outcome_cells[1].text = oc.get('pi', '')
        outcome_cells[2].text = oc.get('baseline', '')
        outcome_cells[3].text = oc.get('target', '')
        outcome_cells[4].text = oc.get('mov', '')
        outcome_cells[5].text = oc.get('assumptions', '')
        for cell in outcome_cells:
            cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F2F2F2"/>'))
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.bold = True
        
        # Now add each output belonging to this outcome
        for op in outputs:
            if op['outcome_index'] == idx:
                op_cells = table.add_row().cells
                op_cells[0].text = f"Output {op.get('number','')}: {op.get('description','')}"
                op_cells[1].text = op.get('pi', '')
                op_cells[2].text = op.get('baseline', '')
                op_cells[3].text = op.get('target', '')
                op_cells[4].text = op.get('mov', '')
                op_cells[5].text = op.get('assumptions', '')
                for cell in op_cells:
                    cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FFFFFF"/>'))
    
    table.autofit = True



def generate_workplan(doc):
    heading = doc.add_heading("WORKPLAN", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    try:
        start_date = parse_date_string(session['general_info']['date_eod'])
        end_date = parse_date_string(session['general_info']['date_nte'])
    except Exception:
        start_date = datetime.today()
        end_date = datetime.today()
    yq_list = get_years_and_quarters(start_date, end_date)
    total_quarters = sum(len(item["Quarters"]) for item in yq_list)
    total_cols = 2 + total_quarters

    # Build a global list of quarters for smooth spectrum coloring
    global_quarters = []
    for item in yq_list:
        for q in item["Quarters"]:
            global_quarters.append((item["Year"], q))
    total_global = len(global_quarters)

    # Create header table (2 rows)
    header_table = doc.add_table(rows=2, cols=total_cols)
    header_table.style = "Table Grid"
    for r in header_table.rows:
        r._tr.get_or_add_trPr().append(parse_xml(r'<w:tblHeader {0} w:val="true"/>'.format(nsdecls('w'))))
    # Merge first two header cells vertically for "Outputs" and "Activities"
    for idx in [0, 1]:
        merged = header_table.cell(0, idx).merge(header_table.cell(1, idx))
        merged.text = "Outputs" if idx == 0 else "Activities"
        for para in merged.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.runs[0].bold = True
        merged._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="CCE5FF"/>'.format(nsdecls('w'))))
    current_col = 2
    for i, item in enumerate(yq_list):
        num_q = len(item["Quarters"])
        top_cell = header_table.cell(0, current_col)
        if num_q > 1:
            for j in range(1, num_q):
                top_cell = top_cell.merge(header_table.cell(0, current_col + j))
        top_cell.text = f"Year {i+1} ({item['Year']})"
        top_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="D0D0D0"/>'.format(nsdecls('w'))))
        for para in top_cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.runs[0].bold = True
        for j, q in enumerate(item["Quarters"]):
            cell = header_table.cell(1, current_col + j)
            cell.text = q
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="E0E0E0"/>'.format(nsdecls('w'))))
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].bold = True
        current_col += num_q

    # Create body table
    body_table = doc.add_table(rows=0, cols=total_cols)
    body_table.style = "Table Grid"

    outcomes = session.get('outcomes', [])
    outputs_list = session.get('outputs', [])
    activities_dict = session.get('activities', {})

    # For each outcome, add a merged row then for each output and its activities below it.
    for oc_idx, oc in enumerate(outcomes):
        # Outcome row spanning all columns
        outcome_row = body_table.add_row()
        outcome_cell = outcome_row.cells[0]
        for c in range(1, total_cols):
            outcome_cell = outcome_cell.merge(body_table.rows[-1].cells[c])
        outcome_cell.text = f"Outcome {oc_idx+1}: {oc.get('description','')}"
        outcome_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="F2FFF2"/>'.format(nsdecls('w'))))
        for para in outcome_cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.runs[0].bold = True

        # Now for each output under this outcome:
        for op in outputs_list:
            if op['outcome_index'] == oc_idx:
                acts = activities_dict.get(str(outputs_list.index(op)), [])
                if not acts:
                    row = body_table.add_row()
                    cells = row.cells
                    cells[0].text = f"Output {op.get('number','')}: {op.get('description','')}"
                    cells[1].text = "(No Activities)"
                    for col in range(2, total_cols):
                        cells[col].text = ""
                else:
                    first_activity = True
                    for a_idx, act in enumerate(acts):
                        row = body_table.add_row()
                        cells = row.cells
                        if first_activity:
                            cells[0].text = f"Output {op.get('number','')}: {op.get('description','')}"
                            # Vertically merge output cell across activity rows with namespace declared
                            cells[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:vMerge {0} w:val="restart"/>'.format(nsdecls('w'))))
                            first_activity = False
                        else:
                            cells[0].text = ""
                            cells[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:vMerge {0}/> '.format(nsdecls('w'))))
                        cells[1].text = f"Activity {op.get('number','')}.{a_idx+1}: {act.get('description','')}"
                        for para in cells[0].paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for para in cells[1].paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        var_col = 2
                        overall_idx = 0
                        for (year, q) in global_quarters:
                            current_val = year * 10 + get_quarter_number(q)
                            sy = act.get('start_year', start_date.year)
                            sq = get_quarter_number(act.get('start_quarter', 'Q1'))
                            ey = act.get('end_year', end_date.year)
                            eq = get_quarter_number(act.get('end_quarter', 'Q4'))
                            start_val = sy * 10 + sq
                            end_val = ey * 10 + eq
                            cell = cells[var_col]
                            if start_val <= current_val <= end_val:
                                fraction = overall_idx / (total_global - 1) if total_global > 1 else 0
                                fill_color = interpolate_color((210,230,245), (225,245,225), fraction)
                                cell.text = ""
                                cell._tc.get_or_add_tcPr().append(parse_xml(
                                    r'<w:shd {0} w:fill="{1}"/>'.format(nsdecls('w'), fill_color)))
                            else:
                                cell.text = ""
                            var_col += 1
                            overall_idx += 1
    body_table.autofit = True

def generate_budget_table(doc):
    heading = doc.add_heading("BUDGET", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    gi = session.get('general_info', {})
    try:
        date_eod = parse_date_string(gi['date_eod'])
        date_nte = parse_date_string(gi['date_nte'])
    except Exception:
        date_eod = datetime.today()
        date_nte = datetime.today()
    num_years = date_nte.year - date_eod.year + 1
    outcomes = session.get('outcomes', [])
    outputs = session.get('outputs', [])
    output_budgets = session.get('output_budgets', {})
    base_cols = 7
    table = doc.add_table(rows=1, cols=base_cols + num_years + 1)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    hdr = ["Outcomes", "Outputs", "FAO Cost Category", "Budget Title", "Budget Unit", "Number of Units", "Unit Cost (USD)"]
    for i, h in enumerate(hdr):
        header_cells[i].text = h
        for para in header_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.runs[0].bold = True
        header_cells[i]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="CCE5FF"/>'.format(nsdecls('w'))))
    for y in range(num_years):
        cell = header_cells[base_cols + y]
        cell.text = f"Year {y+1} ({date_eod.year + y})"
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.runs[0].bold = True
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="B0D0FF"/>'.format(nsdecls('w'))))
    last_cell = header_cells[base_cols + num_years]
    last_cell.text = "Total Cost (USD)"
    for para in last_cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.runs[0].bold = True
    last_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="CCE5FF"/>'.format(nsdecls('w'))))
    current_row_index = 1
    grand_total = 0
    grand_year_sums = [0]*num_years
    for oc_idx, oc in enumerate(outcomes):
        outcome_row_indices = []
        for op_idx, op in enumerate(outputs):
            if op['outcome_index'] == oc_idx:
                b_list = output_budgets.get(str(op_idx), [])
                for b_entry in b_list:
                    new_row = table.add_row()
                    cells = new_row.cells
                    cells[1].text = f"Output {op['number']}: {op['description']}"
                    cells[2].text = b_entry['fao_category']
                    cells[3].text = b_entry['budget_title']
                    cells[4].text = b_entry['budget_unit']
                    cells[5].text = FormatWithHardSpace(b_entry['no_of_units'], ",.0f")
                    cells[6].text = FormatWithHardSpace(b_entry['unit_cost'], ",.0f")
                    y_allocs = b_entry['year_allocations']
                    for y_i in range(num_years):
                        cells[base_cols + y_i].text = FormatWithHardSpace(y_allocs[y_i] if y_i < len(y_allocs) else 0, ",.0f")
                    cells[base_cols + num_years].text = FormatWithHardSpace(b_entry['total_cost'], ",.0f")
                    outcome_row_indices.append(len(table.rows)-1)
                    grand_total += b_entry['total_cost']
                    for y_i in range(num_years):
                        grand_year_sums[y_i] += (y_allocs[y_i] if y_i < len(y_allocs) else 0)
        if outcome_row_indices:
            first_row = outcome_row_indices[0]
            last_row = outcome_row_indices[-1]
            merged_cell = table.cell(first_row, 0).merge(table.cell(last_row, 0))
            merged_cell.text = f"Outcome {oc_idx+1}: {oc.get('description','')}"
        total_row = table.add_row()
        total_cells = total_row.cells
        merged_total = total_cells[0]
        for c in range(1, base_cols):
            merged_total = merged_total.merge(total_cells[c])
        merged_total.text = f"Total Cost for Outcome {oc_idx+1}"
        outcome_total = 0
        outcome_year_sums = [0]*num_years
        for op_idx, op in enumerate(outputs):
            if op['outcome_index'] == oc_idx:
                b_list = output_budgets.get(str(op_idx), [])
                for b_entry in b_list:
                    outcome_total += b_entry['total_cost']
                    for y_i in range(num_years):
                        outcome_year_sums[y_i] += (b_entry['year_allocations'][y_i] if y_i < len(b_entry['year_allocations']) else 0)
        for y_i in range(num_years):
            total_cells[base_cols + y_i].text = FormatWithHardSpace(outcome_year_sums[y_i], ",.0f")
        total_cells[base_cols + num_years].text = FormatWithHardSpace(outcome_total, ",.0f")
        for cell in total_cells:
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="E0E0E0"/>'.format(nsdecls('w'))))
    grand_row = table.add_row()
    grand_cells = grand_row.cells
    merged_grand = grand_cells[0]
    for c in range(1, base_cols):
        merged_grand = merged_grand.merge(grand_cells[c])
    merged_grand.text = "Grand Total for All Outputs"
    for y_i in range(num_years):
        grand_cells[base_cols + y_i].text = FormatWithHardSpace(grand_year_sums[y_i], ",.0f")
    grand_cells[base_cols + num_years].text = FormatWithHardSpace(grand_total, ",.0f")
    for cell in grand_cells:
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="BFBFBF"/>'.format(nsdecls('w'))))
    table.autofit = True

def generate_budget_fao_table(doc):
    heading = doc.add_heading("BUDGET (FAO Cost Category)", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    gi = session.get('general_info', {})
    try:
        date_eod = parse_date_string(gi['date_eod'])
        date_nte = parse_date_string(gi['date_nte'])
    except Exception:
        date_eod = datetime.today()
        date_nte = datetime.today()
    num_years = date_nte.year - date_eod.year + 1
    outputs = session.get('outputs', [])
    output_budgets = session.get('output_budgets', {})
    fao_dict = {}
    for op_idx, op in enumerate(outputs):
        b_list = output_budgets.get(str(op_idx), [])
        for b_entry in b_list:
            cat = b_entry['fao_category']
            fao_dict.setdefault(cat, []).append((op, b_entry))
    base_cols = 6
    table = doc.add_table(rows=1, cols=base_cols + num_years + 1)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    headers = ["FAO Cost Category", "Output", "Budget Title", "Budget Unit", "Number of Units", "Unit Cost (USD)"]
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for para in header_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.runs[0].bold = True
        header_cells[i]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="CCE5FF"/>'.format(nsdecls('w'))))
    for y in range(num_years):
        cell = header_cells[base_cols + y]
        cell.text = f"Year {y+1} ({date_eod.year + y})"
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.runs[0].bold = True
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="B0D0FF"/>'.format(nsdecls('w'))))
    last_cell = header_cells[base_cols + num_years]
    last_cell.text = "Total Cost (USD)"
    for para in last_cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.runs[0].bold = True
    last_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="CCE5FF"/>'.format(nsdecls('w'))))
    current_row_index = 1
    grand_total = 0
    grand_year_sums = [0]*num_years
    for cat, entry_list in fao_dict.items():
        if not entry_list:
            continue
        cat_row_indices = []
        cat_total = 0
        cat_year_sums = [0]*num_years
        for (op, b_entry) in entry_list:
            new_row = table.add_row()
            cells = new_row.cells
            cells[1].text = f"Output {op.get('number','')}"
            cells[2].text = b_entry['budget_title']
            cells[3].text = b_entry['budget_unit']
            cells[4].text = FormatWithHardSpace(b_entry['no_of_units'], ",.0f")
            cells[5].text = FormatWithHardSpace(b_entry['unit_cost'], ",.0f")
            y_allocs = b_entry['year_allocations']
            for y_i in range(num_years):
                cells[base_cols + y_i].text = FormatWithHardSpace(y_allocs[y_i] if y_i < len(y_allocs) else 0, ",.0f")
            cells[base_cols + num_years].text = FormatWithHardSpace(b_entry['total_cost'], ",.0f")
            cat_row_indices.append(len(table.rows)-1)
            cat_total += b_entry['total_cost']
            for y_i in range(num_years):
                cat_year_sums[y_i] += (y_allocs[y_i] if y_i < len(y_allocs) else 0)
                grand_year_sums[y_i] += (y_allocs[y_i] if y_i < len(y_allocs) else 0)
            grand_total += b_entry['total_cost']
        if cat_row_indices:
            first_row = cat_row_indices[0]
            last_row = cat_row_indices[-1]
            merged_cat = table.cell(first_row, 0).merge(table.cell(last_row, 0))
            merged_cat.text = cat
        total_row = table.add_row()
        total_cells = total_row.cells
        merged_total = total_cells[0]
        for c in range(1, base_cols):
            merged_total = merged_total.merge(total_cells[c])
        merged_total.text = f"Total for {cat}"
        for y_i in range(num_years):
            total_cells[base_cols + y_i].text = FormatWithHardSpace(cat_year_sums[y_i], ",.0f")
        total_cells[base_cols + num_years].text = FormatWithHardSpace(cat_total, ",.0f")
        for cell in total_cells:
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="E0E0E0"/>'.format(nsdecls('w'))))
    grand_row = table.add_row()
    grand_cells = grand_row.cells
    merged_grand = grand_cells[0]
    for c in range(1, base_cols):
        merged_grand = merged_grand.merge(grand_cells[c])
    merged_grand.text = "Grand Total for All Categories"
    for y_i in range(num_years):
        grand_cells[base_cols + y_i].text = FormatWithHardSpace(grand_year_sums[y_i], ",.0f")
    grand_cells[base_cols + num_years].text = FormatWithHardSpace(grand_total, ",.0f")
    for cell in grand_cells:
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="BFBFBF"/>'.format(nsdecls('w'))))
    table.autofit = True

def generate_goods_table(doc):
    heading = doc.add_heading("Procurement of Goods", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    goods_list = session.get('goods', [])
    if not goods_list:
        return
    headers = [
        "Project Related Activity",
        "Required Good",
        "Unit of Measure",
        "Estimated Quantities",
        "Unit Price (USD)",
        "Estimated Total Cost (USD)",
        "Procurement Method",
        "Tender Launch Date",
        "Contract Award Date",
        "Delivery Date",
        "Destination Terms",
        "Status",
        "Constraints"
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    row0 = table.rows[0].cells
    for i, h in enumerate(headers):
        row0[i].text = h
        for para in row0[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.runs[0].bold = True
        row0[i]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="CCFFCC"/>'.format(nsdecls('w'))))
    for g in goods_list:
        new_row = table.add_row().cells
        new_row[0].text = g.get('project_activity','')
        new_row[1].text = g.get('required_good','')
        new_row[2].text = g.get('unit_measure','')
        new_row[3].text = str(g.get('estimated_qty',''))
        new_row[4].text = str(g.get('unit_price',''))
        new_row[5].text = str(g.get('estimated_total_cost',''))
        new_row[6].text = g.get('procurement_method','')
        new_row[7].text = g.get('tender_launch_date','')
        new_row[8].text = g.get('contract_award_date','')
        new_row[9].text = g.get('delivery_date','')
        new_row[10].text = g.get('final_destination_terms','')
        new_row[11].text = g.get('status','')
        new_row[12].text = g.get('constraints','')
        for cell in new_row:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True

def generate_subcontracting_table(doc):
    heading = doc.add_heading("Procurement of Commercial Services", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    services_list = session.get('services', [])
    if not services_list:
        return
    headers = [
        "Project Activity",
        "Required Service",
        "Estimated Number of Contracts",
        "Unit Price (USD)",
        "Estimated Total Cost (USD)",
        "Procurement Method",
        "Tender Launch Date",
        "Contract Award Date",
        "Delivery Date",
        "Status",
        "Constraints"
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    row0 = table.rows[0].cells
    for i, h in enumerate(headers):
        row0[i].text = h
        for para in row0[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.runs[0].bold = True
        row0[i]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {0} w:fill="CCFFCC"/>'.format(nsdecls('w'))))
    for s in services_list:
        new_row = table.add_row().cells
        new_row[0].text = s.get('project_activity','')
        new_row[1].text = s.get('required_service','')
        new_row[2].text = str(s.get('estimated_number_of_contracts',''))
        new_row[3].text = str(s.get('unit_price',''))
        new_row[4].text = str(s.get('estimated_total_cost',''))
        new_row[5].text = s.get('procurement_method','')
        new_row[6].text = s.get('tender_launch_date','')
        new_row[7].text = s.get('contract_award_date','')
        new_row[8].text = s.get('delivery_date','')
        new_row[9].text = s.get('status','')
        new_row[10].text = s.get('constraints','')
        for cell in new_row:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True

# -------------------------
# Flask Routes
# -------------------------
@app.route('/')
def index():
    init_session()
    return redirect(url_for('general'))

@app.route('/general', methods=['GET', 'POST'])
def general():
    init_session()
    if request.method == 'POST':
        year_eod = request.form.get('year_eod')
        month_eod = request.form.get('month_eod')
        day_eod = request.form.get('day_eod')
        year_nte = request.form.get('year_nte')
        month_nte = request.form.get('month_nte')
        day_nte = request.form.get('day_nte')
        project_title = request.form.get('project_title')
        project_symbol = request.form.get('project_symbol')
        operational_partner = request.form.get('operational_partner')
        opa_value = request.form.get('opa_value')
        if not (year_eod and month_eod and day_eod and year_nte and month_nte and day_nte and 
                project_title and project_symbol and operational_partner and opa_value):
            flash("All fields are required!")
            return render_template('general_info.html', general_info=request.form)
        try:
            opa_val = float(opa_value)
        except ValueError:
            flash("OPA Value must be numeric!")
            return render_template('general_info.html', general_info=request.form)
        try:
            date_eod = datetime.strptime(f"{year_eod} {month_eod} {day_eod}", "%Y %B %d")
            date_nte = datetime.strptime(f"{year_nte} {month_nte} {day_nte}", "%Y %B %d")
        except ValueError:
            flash("Invalid date. Please check your date fields.")
            return render_template('general_info.html', general_info=request.form)
        if date_eod > date_nte:
            flash("OPA EOD date cannot be after OPA NTE date. Please fix the dates.")
            return render_template('general_info.html', general_info=request.form)
        session['general_info'] = {
            'date_eod': date_eod.strftime("%Y-%m-%d"),
            'year_eod': date_eod.strftime("%Y"),
            'month_eod': date_eod.strftime("%B"),
            'day_eod': str(date_eod.day),
            'date_nte': date_nte.strftime("%Y-%m-%d"),
            'year_nte': date_nte.strftime("%Y"),
            'month_nte': date_nte.strftime("%B"),
            'day_nte': str(date_nte.day),
            'project_title': project_title,
            'project_symbol': project_symbol,
            'operational_partner': operational_partner,
            'opa_value': opa_val
        }
        flash("General Information saved. Proceed to Outcome entry...")
        return redirect(url_for('outcomes'))
    else:
        return render_template('general_info.html', general_info=session.get('general_info', {}))

@app.route('/outcomes', methods=['GET', 'POST'])
def outcomes():
    init_session()
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'back':
            return redirect(url_for('general'))
        elif action == 'delete':
            try:
                index = int(request.form.get('delete_index'))
                outcomes_list = session.get('outcomes', [])
                if 0 <= index < len(outcomes_list):
                    delete_outcome_cascade(index)
                    flash("Outcome and all associated outputs, budgets, and activities deleted.")
                else:
                    flash("Invalid outcome index.")
            except Exception:
                flash("Error deleting outcome.")
            return redirect(url_for('outcomes'))
        elif action == 'edit':
            try:
                index = int(request.form.get('edit_index'))
                return redirect(url_for('edit_outcome', index=index))
            except Exception:
                flash("Invalid outcome index for edit.")
                return redirect(url_for('outcomes'))
        elif action == 'add':
            desc = request.form.get('outcome_desc')
            pi = request.form.get('pi')
            baseline = request.form.get('baseline')
            target = request.form.get('target')
            mov = request.form.get('mov')
            assumptions = request.form.get('assumptions')
            if not (desc and pi and baseline and target and mov and assumptions):
                flash("Please fill in all outcome fields.")
                outcome_objs = [Outcome.from_dict(o) for o in session.get('outcomes', [])]
                return render_template('outcomes.html', outcomes=outcome_objs)
            outcomes_list = session.get('outcomes', [])
            new_outcome = Outcome(desc, pi, baseline, target, mov, assumptions)
            outcomes_list.append(new_outcome.to_dict())
            session['outcomes'] = outcomes_list
            flash("Outcome added!")
            return redirect(url_for('outcomes'))
        elif action == 'next':
            outcomes_list = session.get('outcomes', [])
            if not outcomes_list:
                flash("Please add at least one outcome before proceeding.")
                outcome_objs = [Outcome.from_dict(o) for o in outcomes_list]
                return render_template('outcomes.html', outcomes=outcome_objs)
            return redirect(url_for('outputs'))
    else:
        outcomes_list = session.get('outcomes', [])
        outcome_objs = [Outcome.from_dict(o) for o in outcomes_list]
        return render_template('outcomes.html', outcomes=outcome_objs)

@app.route('/edit_outcome/<int:index>', methods=['GET', 'POST'])
def edit_outcome(index):
    init_session()
    outcomes_list = session.get('outcomes', [])
    if index < 0 or index >= len(outcomes_list):
        flash("Invalid outcome index.")
        return redirect(url_for('outcomes'))
    if request.method == 'POST':
        desc = request.form.get('outcome_desc')
        pi = request.form.get('pi')
        baseline = request.form.get('baseline')
        target = request.form.get('target')
        mov = request.form.get('mov')
        assumptions = request.form.get('assumptions')
        if not (desc and pi and baseline and target and mov and assumptions):
            flash("All fields are required.")
            return render_template('edit_outcome.html', outcome=outcomes_list[index], index=index)
        updated = Outcome(desc, pi, baseline, target, mov, assumptions)
        outcomes_list[index] = updated.to_dict()
        session['outcomes'] = outcomes_list
        flash("Outcome updated.")
        return redirect(url_for('outcomes'))
    else:
        return render_template('edit_outcome.html', outcome=outcomes_list[index], index=index)

@app.route('/outputs', methods=['GET', 'POST'])
def outputs():
    init_session()
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'back':
            return redirect(url_for('outcomes'))
        elif action == 'delete':
            try:
                index = int(request.form.get('delete_index'))
                outputs_list = session.get('outputs', [])
                if 0 <= index < len(outputs_list):
                    delete_output_cascade(index)
                    flash("Output and its associated budgets and activities deleted.")
                else:
                    flash("Invalid output index.")
            except Exception:
                flash("Error deleting output.")
            return redirect(url_for('outputs'))
        elif action == 'edit':
            try:
                index = int(request.form.get('edit_index'))
                return redirect(url_for('edit_output', index=index))
            except Exception:
                flash("Invalid output index for edit.")
                return redirect(url_for('outputs'))
        elif action == 'add':
            try:
                outcome_index = int(request.form.get('selected_outcome'))
            except Exception:
                flash("Invalid outcome selection.")
                return redirect(url_for('outputs'))
            output_desc = request.form.get('output_desc')
            pi = request.form.get('pi')
            baseline = request.form.get('baseline')
            target = request.form.get('target')
            mov = request.form.get('mov')
            assumptions = request.form.get('assumptions')
            if not (output_desc and pi and baseline and target and mov and assumptions):
                flash("Please fill in all output fields.")
                return redirect(url_for('outputs'))
            outputs_list = session.get('outputs', [])
            count = sum(1 for o in outputs_list if o['outcome_index'] == outcome_index)
            new_number = f"{outcome_index + 1}.{count + 1}"
            new_output = Output(outcome_index, output_desc, pi, baseline, target, mov, assumptions, new_number)
            outputs_list.append(new_output.to_dict())
            session['outputs'] = outputs_list
            recalc_output_numbers()
            flash("Output added!")
            return redirect(url_for('outputs'))
        elif action == 'next':
            outputs_list = session.get('outputs', [])
            outcomes_list = session.get('outcomes', [])
            missing = []
            for i in range(len(outcomes_list)):
                if not any(o['outcome_index'] == i for o in outputs_list):
                    missing.append(f"Outcome {i+1}")
            if missing:
                flash("Please add at least one output for " + ", ".join(missing) + " before proceeding.")
                return redirect(url_for('outputs'))
            return redirect(url_for('output_budget'))
    else:
        outcomes_list = session.get('outcomes', [])
        outputs_list = session.get('outputs', [])
        indexed_outputs = [(i, Output.from_dict(o)) for i, o in enumerate(outputs_list)]
        indexed_outputs.sort(key=lambda x: float(x[1].number))
        outputs_by_outcome = {}
        for global_index, output in indexed_outputs:
            outputs_by_outcome.setdefault(output.outcome_index, []).append((global_index, output))
        outcome_objs = [Outcome.from_dict(o) for o in outcomes_list]
        return render_template('outputs.html', outcomes=outcome_objs, outputs_by_outcome=outputs_by_outcome)

@app.route('/edit_output/<int:index>', methods=['GET', 'POST'])
def edit_output(index):
    init_session()
    outputs_list = session.get('outputs', [])
    if index < 0 or index >= len(outputs_list):
        flash("Invalid output index.")
        return redirect(url_for('outputs'))
    current_output = Output.from_dict(outputs_list[index])
    outcomes_list = session.get('outcomes', [])
    associated_outcome = {}
    if current_output.outcome_index < len(outcomes_list):
        associated_outcome = outcomes_list[current_output.outcome_index]
    if request.method == 'POST':
        output_desc = request.form.get('output_desc')
        pi = request.form.get('pi')
        baseline = request.form.get('baseline')
        target = request.form.get('target')
        mov = request.form.get('mov')
        assumptions = request.form.get('assumptions')
        if not (output_desc and pi and baseline and target and mov and assumptions):
            flash("All fields are required.")
            return render_template('edit_output.html', output=outputs_list[index], index=index, outcome=associated_outcome)
        updated_output = Output(current_output.outcome_index, output_desc, pi, baseline, target, mov, assumptions, current_output.number)
        outputs_list[index] = updated_output.to_dict()
        session['outputs'] = outputs_list
        recalc_output_numbers()
        flash("Output updated.")
        return redirect(url_for('outputs'))
    else:
        return render_template('edit_output.html', output=outputs_list[index], index=index, outcome=associated_outcome)

@app.route('/output_budget', methods=['GET', 'POST'])
def output_budget():
    init_session()
    gi = session.get('general_info', {})
    try:
        date_eod = datetime.strptime(gi['date_eod'], "%Y-%m-%d")
        date_nte = datetime.strptime(gi['date_nte'], "%Y-%m-%d")
    except (KeyError, ValueError):
        flash("General info incomplete.")
        return redirect(url_for('general'))
    num_years = date_nte.year - date_eod.year + 1
    if request.method == 'POST':
        outputs_list = session.get('outputs', [])
        output_budgets = {}
        for i, out in enumerate(outputs_list):
            i_str = str(i)
            budget_entries = []
            j = 0
            while True:
                base_key = f'fao_{i}_{j}'
                if base_key not in request.form:
                    break
                fao = request.form.get(f'fao_{i}_{j}', '')
                title = request.form.get(f'title_{i}_{j}', '')
                unit = request.form.get(f'unit_{i}_{j}', '')
                no_units_str = request.form.get(f'no_units_{i}_{j}', '0')
                unit_cost_str = request.form.get(f'unit_cost_{i}_{j}', '0')
                year_allocs = []
                for k in range(1, num_years + 1):
                    val_str = request.form.get(f'year_{i}_{j}_{k}', '0')
                    try:
                        val = float(val_str)
                    except ValueError:
                        val = 0
                    year_allocs.append(val)
                try:
                    no_units_val = float(no_units_str)
                except ValueError:
                    no_units_val = 0
                try:
                    unit_cost_val = float(unit_cost_str)
                except ValueError:
                    unit_cost_val = 0
                total_cost = no_units_val * unit_cost_val
                budget_entry = {
                    'fao_category': fao,
                    'budget_title': title,
                    'budget_unit': unit,
                    'no_of_units': no_units_val,
                    'unit_cost': unit_cost_val,
                    'total_cost': total_cost,
                    'year_allocations': year_allocs
                }
                budget_entries.append(budget_entry)
                j += 1
            output_budgets[str(i)] = budget_entries
        session['output_budgets'] = output_budgets
        flash("Budgets saved. Proceed to Activities.")
        return redirect(url_for('activities'))
    else:
        outputs_list = session.get('outputs', [])
        output_budgets = session.get('output_budgets', {})
        for i, out in enumerate(outputs_list):
            if str(i) not in output_budgets:
                output_budgets[str(i)] = []
        session['output_budgets'] = output_budgets
        outputs_by_outcome = {}
        for i, out in enumerate(outputs_list):
            idx = out['outcome_index']
            outputs_by_outcome.setdefault(idx, []).append((i, out))
        return render_template('output_budget.html',
                               outputs_by_outcome=outputs_by_outcome,
                               num_years=num_years,
                               outcomes=session.get('outcomes', []),
                               general_info=gi,
                               output_budgets=output_budgets)


@app.route('/activities', methods=['GET', 'POST'])
def activities():
    init_session()
    general_info = session.get('general_info', {})
    try:
        opa_eod_year = int(general_info.get('year_eod'))
        opa_nte_year = int(general_info.get('year_nte'))
    except (TypeError, ValueError):
        flash("General information is incomplete. Please re-enter the general info.")
        return redirect(url_for('general'))

    # Compute allowed years and quarters (assume you already have quarter_of_month defined)
    month_map = {
        "January": 1, "February": 2, "March": 3, "April": 4,
        "May": 5, "June": 6, "July": 7, "August": 8,
        "September": 9, "October": 10, "November": 11, "December": 12
    }
    opa_eod_month = month_map.get(general_info.get('month_eod', 'January'), 1)
    opa_nte_month = month_map.get(general_info.get('month_nte', 'December'), 12)
    opa_eod_quarter = quarter_of_month(opa_eod_month)
    opa_nte_quarter = quarter_of_month(opa_nte_month)
    years = list(range(opa_eod_year, opa_nte_year + 1))
    partial_quarters_map_start = {}
    partial_quarters_map_end = {}
    for y in years:
        if y == opa_eod_year:
            partial_quarters_map_start[y] = ["Q" + str(q) for q in range(opa_eod_quarter, 5)]
        else:
            partial_quarters_map_start[y] = ["Q1", "Q2", "Q3", "Q4"]
        if y == opa_nte_year:
            partial_quarters_map_end[y] = ["Q" + str(q) for q in range(1, opa_nte_quarter + 1)]
        else:
            partial_quarters_map_end[y] = ["Q1", "Q2", "Q3", "Q4"]
    partial_quarters_map = {"start": partial_quarters_map_start, "end": partial_quarters_map_end}

    if request.method == 'POST':
        action = request.form.get('action')
        # If the back button was clicked, go to output_budget page
        if action == 'back':
            return redirect(url_for('output_budget'))

        # Otherwise, process the submitted activities...
        outputs_list = session.get('outputs', [])
        activities_dict = {}
        error_flag = False

        for i, out in enumerate(outputs_list):
            activities_list = []
            j = 0
            while True:
                s_year_str = request.form.get(f'start_year_{i}_{j}', '').strip()
                if s_year_str == '':
                    break
                s_quarter_str = request.form.get(f'start_quarter_{i}_{j}', '').strip()
                e_year_str = request.form.get(f'end_year_{i}_{j}', '').strip()
                e_quarter_str = request.form.get(f'end_quarter_{i}_{j}', '').strip()
                act_desc = request.form.get(f'activity_desc_{i}_{j}', '').strip()

                if s_year_str and s_quarter_str and e_year_str and e_quarter_str and act_desc:
                    try:
                        activity = {
                            'description': act_desc,
                            'start_year': int(s_year_str),
                            'start_quarter': s_quarter_str,
                            'end_year': int(e_year_str),
                            'end_quarter': e_quarter_str
                        }
                    except ValueError:
                        flash("Invalid numeric values in activity dates.")
                        return redirect(url_for('activities'))
                    activities_list.append(activity)
                else:
                    flash("Please fill in all fields for each activity.")
                    error_flag = True
                    break
                j += 1

            if not activities_list:
                flash(f"Output {i+1} must have at least one activity.")
                error_flag = True
            activities_dict[str(i)] = activities_list

        if error_flag:
            return redirect(url_for('activities'))

        session['activities'] = activities_dict
        flash("Activities saved!")
        return redirect(url_for('procurement'))
    else:
        outputs_list = session.get('outputs', [])
        activities_dict = session.get('activities', {})
        return render_template('activities.html',
                               outputs=outputs_list,
                               activities_dict=activities_dict,
                               years=years,
                               partialQuartersMap=partial_quarters_map)


@app.route('/procurement', methods=['GET', 'POST'])
def procurement():
    init_session()
    if request.method == 'POST':
        action = request.form.get('action', '').lower()
        if action == 'select_types':
            use_goods = 'use_goods' in request.form
            use_services = 'use_services' in request.form
            session['use_goods'] = use_goods
            session['use_services'] = use_services
            return redirect(url_for('procurement'))
        elif action == 'add_goods':
            try:
                dt_tender = datetime(int(request.form.get('pog_tender_launch_year')),
                                     list(calendar.month_name).index(request.form.get('pog_tender_launch_month')),
                                     int(request.form.get('pog_tender_launch_day')))
                dt_award = datetime(int(request.form.get('pog_contract_award_year')),
                                    list(calendar.month_name).index(request.form.get('pog_contract_award_month')),
                                    int(request.form.get('pog_contract_award_day')))
                dt_delivery = datetime(int(request.form.get('pog_delivery_year')),
                                       list(calendar.month_name).index(request.form.get('pog_delivery_month')),
                                       int(request.form.get('pog_delivery_day')))
            except ValueError:
                flash("Invalid date for Procurement of Goods. Please check the day/month/year.")
                return redirect(url_for('procurement'))
            if not (dt_tender <= dt_award <= dt_delivery):
                flash("Tender Launch, Contract Award, and Delivery dates must be in order (Tender ≤ Award ≤ Delivery).")
                return redirect(url_for('procurement'))
            row = {
                "project_activity": request.form.get("pog_project_activity", "").strip(),
                "required_good": request.form.get("pog_required_good", "").strip(),
                "unit_measure": request.form.get("pog_unit_measure", "").strip(),
                "estimated_qty": request.form.get("pog_estimated_qty", "").strip(),
                "unit_price": request.form.get("pog_unit_price", "").strip(),
                "estimated_total_cost": request.form.get("pog_estimated_total_cost", "").strip(),
                "procurement_method": request.form.get("pog_procurement_method", "").strip(),
                "tender_launch_date": dt_tender.strftime("%Y-%m-%d"),
                "contract_award_date": dt_award.strftime("%Y-%m-%d"),
                "delivery_date": dt_delivery.strftime("%Y-%m-%d"),
                "final_destination_terms": request.form.get("pog_final_destination_terms", "").strip(),
                "status": request.form.get("pog_status", "").strip(),
                "constraints": request.form.get("pog_constraints", "").strip()
            }
            session['goods'].append(row)
            flash("Procurement of Goods entry added.")
            return redirect(url_for('procurement'))
        elif action == 'delete_goods':
            try:
                idx = int(request.form.get('delete_goods_index', ''))
                goods_list = session['goods']
                if 0 <= idx < len(goods_list):
                    goods_list.pop(idx)
                    session['goods'] = goods_list
                    flash("Procurement of Goods entry deleted.")
                else:
                    flash("Invalid goods index to delete.")
            except ValueError:
                flash("Invalid goods index to delete.")
            return redirect(url_for('procurement'))
        elif action == 'add_services':
            try:
                dt_tender = datetime(int(request.form.get('scs_tender_launch_year')),
                                     list(calendar.month_name).index(request.form.get('scs_tender_launch_month')),
                                     int(request.form.get('scs_tender_launch_day')))
                dt_award = datetime(int(request.form.get('scs_contract_award_year')),
                                    list(calendar.month_name).index(request.form.get('scs_contract_award_month')),
                                    int(request.form.get('scs_contract_award_day')))
                dt_delivery = datetime(int(request.form.get('scs_delivery_year')),
                                       list(calendar.month_name).index(request.form.get('scs_delivery_month')),
                                       int(request.form.get('scs_delivery_day')))
            except ValueError:
                flash("Invalid date for Subcontracting of Commercial Services. Please check the day/month/year.")
                return redirect(url_for('procurement'))
            if not (dt_tender <= dt_award <= dt_delivery):
                flash("Tender Launch, Contract Award, and Delivery dates must be in order for Services.")
                return redirect(url_for('procurement'))
            row = {
                "project_activity": request.form.get("scs_project_activity", "").strip(),
                "required_service": request.form.get("scs_required_service", "").strip(),
                "estimated_number_of_contracts": request.form.get("scs_estimated_number_of_contracts", "").strip(),
                "unit_price": request.form.get("scs_unit_price", "").strip(),
                "estimated_total_cost": request.form.get("scs_estimated_total_cost", "").strip(),
                "procurement_method": request.form.get("scs_procurement_method", "").strip(),
                "tender_launch_date": dt_tender.strftime("%Y-%m-%d"),
                "contract_award_date": dt_award.strftime("%Y-%m-%d"),
                "delivery_date": dt_delivery.strftime("%Y-%m-%d"),
                "status": request.form.get("scs_status", "").strip(),
                "constraints": request.form.get("scs_constraints", "").strip()
            }
            session['services'].append(row)
            flash("Subcontracting Services entry added.")
            return redirect(url_for('procurement'))
        elif action == 'delete_services':
            try:
                idx = int(request.form.get('delete_services_index', ''))
                services_list = session['services']
                if 0 <= idx < len(services_list):
                    services_list.pop(idx)
                    session['services'] = services_list
                    flash("Subcontracting Services entry deleted.")
                else:
                    flash("Invalid services index to delete.")
            except ValueError:
                flash("Invalid services index to delete.")
            return redirect(url_for('procurement'))
        elif action in ['finalize', 'finalize_procurement', 'proceed']:
            use_goods = session.get('use_goods', False)
            use_services = session.get('use_services', False)
            if use_goods and len(session['goods']) == 0:
                flash("You indicated procurement of goods, but have not added any. Please add or uncheck.")
                return redirect(url_for('procurement'))
            if use_services and len(session['services']) == 0:
                flash("You indicated subcontracting of services, but have not added any. Please add or uncheck.")
                return redirect(url_for('procurement'))
            flash("Procurement finalized. Now you can review all data before generating tables.")
            return redirect(url_for('review'))
        else:
            flash("Unknown action in procurement.")
            return redirect(url_for('procurement'))
    else:
        use_goods = session.get('use_goods', False)
        use_services = session.get('use_services', False)
        goods_list = session.get('goods', [])
        services_list = session.get('services', [])
        years = list(range(2020, 2051))  # Updated year range 2020 to 2050
        months = list(calendar.month_name)[1:]
        days = list(range(1, 32))
        return render_template('procurement.html',
                               use_goods=use_goods,
                               use_services=use_services,
                               goods=goods_list,
                               services=services_list,
                               years=years,
                               months=months,
                               days=days,
                               back_url=url_for('activities'))




@app.route('/review', methods=['GET', 'POST'])
def review():
    init_session()
    data = {
        'general_info': session.get('general_info', {}),
        'outcomes': session.get('outcomes', []),
        'outputs': session.get('outputs', []),
        'output_budgets': session.get('output_budgets', {}),
        'activities': session.get('activities', {}),
        'goods': session.get('goods', []),
        'services': session.get('services', []),
        'use_goods': session.get('use_goods', False),
        'use_services': session.get('use_services', False)
    }
    if request.method == 'POST':
        return redirect(url_for('finalize'))
    return render_template('review.html', data=data)

@app.route('/finalize')
def finalize():
    doc = generate_word_doc()
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    filename = "OPA_Tables.docx"
    return send_file(f, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route('/clear')
def clear():
    session.clear()
    flash("Session cleared.")
    return redirect(url_for('general'))

@app.route('/delete_output/<int:index>', methods=['POST'])
def delete_output(index):
    init_session()
    delete_output_cascade(index)
    flash("Output and its associated budgets and activities deleted.")
    return redirect(url_for('outputs'))

@app.route('/generate_doc')
def generate_doc():
    # Replace this with your document-generation logic.
    return "Word document generated (placeholder)."


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))  # Get the assigned port
    app.run(host="0.0.0.0", port=port, debug=True)  # Allow external access
