from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_word_doc(session):
    doc = Document()
    # --- HEADER ---
    gi = session.get('general_info', {})
    header_text = (
        f"Project Symbol: {gi.get('project_symbol', 'N/A')}\n"
        f"Operational Partner: {gi.get('operational_partner', 'N/A')}\n"
        f"Project Title: {gi.get('project_title', 'N/A')}\n"
        f"OPA Value: {gi.get('opa_value', 'N/A')}\n"
    )
    p = doc.add_paragraph(header_text)
    p.runs[0].font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("\n")
    
    # --- 1. RESULTS MATRIX ---
    outcomes = session.get('outcomes', [])
    outputs = session.get('outputs', [])
    if outcomes:
        doc.add_heading("RESULTS MATRIX", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Result Statement", "Performance Indicators", "Baseline", "Target", "Means of Verification", "Assumptions"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        # Add outcome rows and then output rows
        for idx, outcome in enumerate(outcomes):
            row_cells = table.add_row().cells
            row_cells[0].text = outcome.get('description', '')
            row_cells[1].text = outcome.get('pi', '')
            row_cells[2].text = outcome.get('baseline', '')
            row_cells[3].text = outcome.get('target', '')
            row_cells[4].text = outcome.get('mov', '')
            row_cells[5].text = outcome.get('assumptions', '')
            for output in outputs:
                if output.get('outcome_index') == idx:
                    row_cells = table.add_row().cells
                    row_cells[0].text = "Output " + output.get('number', '') + ": " + output.get('description', '')
                    row_cells[1].text = output.get('pi', '')
                    row_cells[2].text = output.get('baseline', '')
                    row_cells[3].text = output.get('target', '')
                    row_cells[4].text = output.get('mov', '')
                    row_cells[5].text = output.get('assumptions', '')
        doc.add_paragraph("\n")
    
    # --- 2. WORKPLAN ---
    # (A simplified version: listing outputs and their activities with timeline info)
    doc.add_heading("WORKPLAN", level=2)
    workplan_table = doc.add_table(rows=1, cols=5)
    workplan_table.style = 'Table Grid'
    wp_hdr = workplan_table.rows[0].cells
    wp_hdr[0].text = "Outputs"
    wp_hdr[1].text = "Activities"
    wp_hdr[2].text = "Start"
    wp_hdr[3].text = "End"
    wp_hdr[4].text = "Timeline"
    activities = session.get('activities', {})
    for i, output in enumerate(outputs):
        acts = activities.get(str(i), [])
        if not acts:
            acts = [{"description": "", "start_year": "", "start_quarter": "", "end_year": "", "end_quarter": ""}]
        for act in acts:
            row_cells = workplan_table.add_row().cells
            row_cells[0].text = "Output " + output.get('number', '') + ": " + output.get('description', '')
            row_cells[1].text = act.get('description', '')
            row_cells[2].text = f"{act.get('start_year', '')} {act.get('start_quarter', '')}"
            row_cells[3].text = f"{act.get('end_year', '')} {act.get('end_quarter', '')}"
            row_cells[4].text = f"{act.get('start_year', '')} {act.get('start_quarter', '')} to {act.get('end_year', '')} {act.get('end_quarter', '')}"
    doc.add_paragraph("\n")
    
    # --- 3. BUDGET ---
    output_budgets = session.get('output_budgets', {})
    if output_budgets:
        doc.add_heading("BUDGET", level=2)
        for out_idx, budgets in output_budgets.items():
            doc.add_paragraph("Output " + str(int(out_idx)+1) + " Budgets:")
            if budgets:
                budget_table = doc.add_table(rows=1, cols=5)
                budget_table.style = 'Table Grid'
                b_hdr = budget_table.rows[0].cells
                b_hdr[0].text = "FAO Category"
                b_hdr[1].text = "Budget Title"
                b_hdr[2].text = "Budget Unit"
                b_hdr[3].text = "No. of Units"
                b_hdr[4].text = "Total Cost"
                for entry in budgets:
                    row_cells = budget_table.add_row().cells
                    row_cells[0].text = entry.get('fao_category', '')
                    row_cells[1].text = entry.get('budget_title', '')
                    row_cells[2].text = entry.get('budget_unit', '')
                    row_cells[3].text = str(entry.get('no_of_units', ''))
                    row_cells[4].text = str(entry.get('total_cost', ''))
            doc.add_paragraph("\n")
    
    # --- 4. PROCUREMENT OF GOODS ---
    goods = session.get('goods', [])
    if goods:
        doc.add_heading("Procurement of Goods", level=2)
        goods_table = doc.add_table(rows=1, cols=13)
        goods_table.style = 'Table Grid'
        g_hdr = goods_table.rows[0].cells
        headers = ["Project Activity", "Required Good", "Unit Measure",
                   "Estimated Qty", "Unit Price", "Total Cost", "Procurement Method",
                   "Tender Launch Date", "Contract Award Date", "Delivery Date",
                   "Destination Terms", "Status", "Constraints"]
        for i, h in enumerate(headers):
            g_hdr[i].text = h
            for p in g_hdr[i].paragraphs:
                p.runs[0].font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for g in goods:
            row_cells = goods_table.add_row().cells
            row_cells[0].text = g.get("project_activity", "")
            row_cells[1].text = g.get("required_good", "")
            row_cells[2].text = g.get("unit_measure", "")
            row_cells[3].text = g.get("estimated_qty", "")
            row_cells[4].text = g.get("unit_price", "")
            row_cells[5].text = g.get("estimated_total_cost", "")
            row_cells[6].text = g.get("procurement_method", "")
            row_cells[7].text = g.get("tender_launch_date", "")
            row_cells[8].text = g.get("contract_award_date", "")
            row_cells[9].text = g.get("delivery_date", "")
            row_cells[10].text = g.get("final_destination_terms", "")
            row_cells[11].text = g.get("status", "")
            row_cells[12].text = g.get("constraints", "")
        doc.add_paragraph("\n")
    
    # --- 5. Subcontracting of Commercial Services ---
    services = session.get('services', [])
    if services:
        doc.add_heading("Subcontracting of Commercial Services", level=2)
        serv_table = doc.add_table(rows=1, cols=10)
        serv_table.style = 'Table Grid'
        s_hdr = serv_table.rows[0].cells
        s_hdr[0].text = "Project Activity"
        s_hdr[1].text = "Required Service"
        s_hdr[2].text = "Estimated Contracts"
        s_hdr[3].text = "Unit Price"
        s_hdr[4].text = "Total Cost"
        s_hdr[5].text = "Procurement Method"
        s_hdr[6].text = "Tender Launch Date"
        s_hdr[7].text = "Contract Award Date"
        s_hdr[8].text = "Delivery Date"
        s_hdr[9].text = "Status"
        for s in services:
            row_cells = serv_table.add_row().cells
            row_cells[0].text = s.get("project_activity", "")
            row_cells[1].text = s.get("required_service", "")
            row_cells[2].text = s.get("estimated_number_of_contracts", "")
            row_cells[3].text = s.get("unit_price", "")
            row_cells[4].text = s.get("estimated_total_cost", "")
            row_cells[5].text = s.get("procurement_method", "")
            row_cells[6].text = s.get("tender_launch_date", "")
            row_cells[7].text = s.get("contract_award_date", "")
            row_cells[8].text = s.get("delivery_date", "")
            row_cells[9].text = s.get("status", "")
        doc.add_paragraph("\n")
    
    return doc
